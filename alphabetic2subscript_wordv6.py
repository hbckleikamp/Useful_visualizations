# -*- coding: utf-8 -*-
"""
Created on Fri Feb 23 17:33:56 2024

@author: e_kle
"""

# import docx2txt
# my_text = docx2txt.process("C:/Users/e_kle/Desktop/Aerobic methanogenesis/SV_ Revisions Selenium Paper/042124 selenium_biorxiv revisions_MVK -test.docx")
# print(my_text)



from docx import Document
import pandas as pd
import numpy as np
#%%
file="C:/Users/e_kle/Desktop/Aerobic methanogenesis/SV_ Revisions Selenium Paper/042124 selenium_biorxiv revisions_MVK -test.docx"
file="C:/Users/e_kle/Desktop/Aerobic methanogenesis/SV_ Revisions Selenium Paper/042124 selenium_biorxiv revisions_MVK -test.docx"
file="C:/Users/e_kle/Desktop/Aerobic methanogenesis/SV_ Revisions Selenium Paper/090224 selenium_paper_no_markdown_ESnT-test.docx"

document = Document(file)


#%% need to somehow separate the runs into ones with distinct styles
font_options=['all_caps','bold','complex_script','cs_bold','cs_italic','double_strike',
'emboss','hidden','imprint','italic','math','name','no_proof',
'outline','rtl','shadow', 'size','small_caps','snap_to_grid','spec_vanish','strike',
'subscript','superscript','underline','web_hidden']

par_format_options=[ 'alignment','first_line_indent',
 'keep_together', 'keep_with_next', 'left_indent', 'line_spacing',
 'line_spacing_rule', 'page_break_before',  'right_indent',
 'space_after', 'space_before',  'widow_control']
#'tab_stops',
#'element'
#'part'
#,'highlight_color'

#need to unpack all paragraph style features
#all run style features

pars=document.paragraphs

parstyle=[]
runstyle=[]
r_ix=0

all_runs=[]
for ip,par in enumerate(pars):
    
     parstyle.append([ip]+[getattr(par.paragraph_format,po) for po in par_format_options])
     
     runs=par.runs
     for run in runs:
         c=getattr(run.font.color,"rgb")
         fa=[getattr(run.font,fo)  for fo in font_options]
         runstyle.append([r_ix,ip]+fa+[c,run.text])                   
         r_ix+=1
         all_runs.append(run)

parstyle=pd.DataFrame(parstyle,columns=["ix"]+par_format_options)
runstyle=pd.DataFrame(runstyle,columns=["run_ix","par_ix"]+font_options+["color","text"])
#runstyle["style"]=runstyle[font_options].notnull().any(axis=1)

merged=[]
t=""

for n,g in runstyle.groupby("par_ix"):
    
    fos=[]
    
    c=0
    for _,r in g.iterrows():
        
        if c:
            if not (fos[-1][font_options].values==r[font_options].values).all():
                d=fos[-1]
                d.text=t
                merged.append(d)
                t=""
        t+=r.text

                
                
        fos.append(r) 
        c+=1
        
    if len(t):
        d=r
        d.text=t
        merged.append(d)
    t=""
        

reftexts=pd.DataFrame([[m.text,m.run_ix] for m in merged],columns=["text","run_ix"])

mdf=pd.DataFrame(merged,columns=runstyle.columns)

#%% get reference order

ms=[]
for mr in merged:
    
    t=mr.text
    
    if "(" in t:
     
        
        
        st=t.replace("(",")").split(")")
        
        rfs=pd.DataFrame(st[1::2],columns=["NY"])
        
        
        rfs["NY"]=rfs["NY"].str.split(",")
        rfs=rfs.explode("NY")
        rfs["NY"]=rfs["NY"].str.strip()
        
        rfs=rfs[rfs.NY!=""]
        if len(rfs):
        
        
            rfs["Name"]=rfs["NY"].str.split().apply(lambda x: " ".join(x[:-1]))
            rfs["Year"]=rfs["NY"].str.split().apply(lambda x: x[-1])
            rfs["run_ix"]=mr.run_ix
            ms.append(rfs)
        
        

    if "References" in t:
        break

   

ms=pd.concat(ms)
ms=ms[(ms["Year"].str.startswith("20")) | (ms["Year"].str.startswith("19"))]
ms.Name=ms.Name.replace("",np.nan).ffill()
ms["m"]=ms["Name"].str.strip()+" "+ms["Year"].str.strip()

ms.index=np.arange(1,len(ms)+1)
    



#%% Parse references list
rfl=[]
tog=0
for par in pars:
    
    t=par.text
    if tog:
        if t=="Supporting information":
            break
        
        rfl.append(t)

    if t=="References":
        tog=1

#merge consecutive rows
c,m,s=0,[],""

while True:
    s+=rfl[c]
    if "(20" in s or "(19" in s: 
        m.append(s)
        s="" 
    c+=1    
    if c>=len(rfl):
        break

m=pd.DataFrame(m,columns=["Ref"])
#rfl=pd.DataFrame([i for i in rfl if i!=""],columns=["Ref"])
m["Name"]=m.Ref.str.split(",").apply(lambda x: x[0])
m["Year"]=m["Ref"].str.replace(")","(",regex=False).str.split("(",regex=False).apply(lambda x: x[1])
m["Year"]=m["Year"].str.split(",").apply(lambda x: x[0]).str.strip()
m["m"]=m["Name"].str.strip()+" "+m["Year"].str.strip()


omms=ms.merge(m[["Ref","m"]],on="m",how="outer")#.drop_duplicates()
mms=ms.reset_index(names=["ap_ix"]).merge(m[["Ref","m"]],on="m",how="left").sort_values(by="ap_ix")#.drop_duplicates()



wrongref=omms[omms.isnull().any(axis=1)]#.sort_values(by="m")

if len(wrongref):
    print("fix incorrect references!!!")
    ""+1

#mms=mms.groupby("Ref",sort=False).nth(0).reset_index(drop=True)


ref_index=mms["Ref"].to_frame().drop_duplicates().reset_index(drop=True).reset_index(names=["Ref_index"])
mms=mms.merge(ref_index,on="Ref",how="left") 
mms=mms.sort_values(by="run_ix")

mms["Ref_index"]+=1
#mms=mms.groupby(["Ref","par_ix"],sort=False).nth(0)

sorted_refs=mms.groupby("Ref_index")["Ref"].nth(0).reset_index(drop=True)
sorted_refs=(sorted_refs.index+1).astype(str)+" "+sorted_refs

rfi=mdf[mdf.text=="References"].index[0]+1

#mdf.loc[rfi:rfi+len(sorted_refs),"text"]

#%%
# c=0
# for n,r in mdf.loc[rfi:].iterrows():
    
#     if r.text in mms["Ref"].tolist():
#         mdf.loc[n,"text"]=sorted_refs[c]

#         c+=1
#     if c==len(sorted_refs):
#         break

#%%

#https://stackoverflow.com/questions/62451416/font-type-and-font-size-in-python-docx
#https://stackoverflow.com/questions/63288920/issues-with-font-and-table-python-docx-module


#%% replace refs with subscripted index


cdoc=Document()
r_ix=-1

for par_ix,par in enumerate(pars):
    
    pt=par.text

    
    p=cdoc.add_paragraph("")
    for po in par_format_options:
        setattr(p.paragraph_format,po,getattr(par.paragraph_format,po))
    
    runs=mdf[mdf.par_ix==par_ix]
    if len(runs):
        
        for _,r in runs.iterrows():
            
            c=0
            run=all_runs[r.run_ix]
            t=r.text
            
            if t in mms["Ref"].tolist():
                break
            
            # if "(Zheng 2018)" in t:
            #     ""+1
            
            rf=mms[mms["run_ix"]==r.run_ix].sort_index()
            if len(rf):
                
                
                st=t.replace("(",")").split(")")
                
                for six,s in enumerate(st):
               
                    if six%2:
                        found=[]
                        for ss in s.split(","):
                            

                            if c<len(rf):
                                if rf["NY"].iloc[c]==ss.strip():
                                    found.append(str(rf["Ref_index"].iloc[c]))
                                    c+=1
                                    
                        if not len(found):
                            d=p.add_run("("+s+")")
                            for fo in font_options:
                                setattr(d.font,fo,getattr(run.font,fo))
                            setattr(d.font.color,"rgb",getattr(run.font.color,"rgb"))
                            
                        else:
                            found=sorted(found)
                            d=p.add_run(",".join(found))
                            d.font.subscript = True
                    
                    else:
                        d=p.add_run(s)
                     
            
            
            else:
                d = p.add_run(t)
                for fo in font_options:
                    setattr(d.font,fo,getattr(run.font,fo))
                setattr(d.font.color,"rgb",getattr(run.font.color,"rgb"))

        

d=p.add_run("\n\n".join(sorted_refs))    
for fo in font_options:
    setattr(d.font,fo,getattr(run.font,fo))
setattr(d.font.color,"rgb",getattr(run.font.color,"rgb"))
#%% write newly structured references    

      
file="test.docx"    
with open(file,"w"): pass #clears file
cdoc.save(file)

    


#%% Update sorted reference list


#%%

#document.save('test2.docx')

#%%
#add paragraph add run add paragraph ...

# for par in pars:
#     print(par.text)

# from docx import *

# document = Document('/path/to/file.docx')
# def bold(document):
#     for para in document.paragraphs:
#         Listbolds = []
#         for run in para.runs:
#             if run.bold:
#                 print(run.text
#                 word = run.text
#                 Listbolds.append(word)
#     return Listbolds